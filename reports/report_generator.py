"""
report_generator.py

Generates a downloadable Microsoft Word (.docx) report from the AI review
results for a given Pull Request.

Public API
----------
parse_review_issues(review_text) -> list[dict]
    Parses the raw AI review string into a list of structured issue dicts.

generate_word_report(repo_name, pr_number, review_text, report_dir) -> str
    Builds the .docx file and returns its absolute path on disk.
"""

import os
import re
import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Internal constants
# ---------------------------------------------------------------------------

# Known field labels that follow the issue description in the LLM output.
# Used to cleanly strip everything after the issue text regardless of whether
# the LLM emits a newline before the label or concatenates them inline.
_NEXT_FIELD_PATTERN = re.compile(
    r"\s*(?:File|Line|Code|Reason|Suggestion)\s*:",
    re.IGNORECASE
)

# Horizontal rule text used inside the .docx
_SEPARATOR = "\u2500" * 40


# ---------------------------------------------------------------------------
# Public: parse_review_issues
# ---------------------------------------------------------------------------

def parse_review_issues(review_text: str) -> list[dict]:
    """
    Parse the raw AI review text into a list of structured issue dicts.

    Each dict has the keys:
        file  (str) – file/line hint extracted from File: or Code: field, or ""
        line  (str) – line reference extracted from Line: or Code: field, or ""
        issue (str) – the cleaned issue description text

    Strategy: split on every 'Issue:' label rather than relying on newlines
    between blocks. This handles both newline-separated and inline-concatenated
    LLM output formats:

        # Newline-separated (standard):
        Issue: Something wrong\nFile: a.py\nLine: 10\nCode: x = y\nReason: ...

        # Inline/concatenated (seen in practice):
        Issue: Something wrongFile: a.pyLine: 10Code: x = yReason: ...

    Fallback: if no 'Issue:' label is found, splits on numbered list markers
    (e.g. '1.', '2.') so the report is never empty.
    """
    issues = []

    # Split on every occurrence of 'Issue:' (case-insensitive).
    # Element [0] is the preamble before the first issue — always skipped.
    raw_blocks = re.split(r"Issue\s*:", review_text, flags=re.IGNORECASE)

    if len(raw_blocks) > 1:
        for raw_block in raw_blocks[1:]:
            raw_block = raw_block.strip()
            if not raw_block:
                continue

            # Extract issue text — stop at the first following field label.
            field_match = _NEXT_FIELD_PATTERN.search(raw_block)
            if field_match:
                issue_text = raw_block[: field_match.start()].strip()
            else:
                issue_text = raw_block.strip()

            # Collapse internal newlines / extra whitespace to a single space.
            issue_text = re.sub(r"\s+", " ", issue_text).strip()

            if not issue_text:
                continue

            # Extract File: and Line: fields if provided explicitly by LLM
            file_match = re.search(
                r"File\s*:\s*(.+?)(?=\s*(?:Line|Code|Reason|Suggestion|Issue)\s*:|$)",
                raw_block,
                re.DOTALL | re.IGNORECASE
            )
            line_match = re.search(
                r"Line\s*:\s*(.+?)(?=\s*(?:File|Code|Reason|Suggestion|Issue)\s*:|$)",
                raw_block,
                re.DOTALL | re.IGNORECASE
            )

            file_hint = (
                file_match.group(1).strip().splitlines()[0].strip()
                if file_match
                else ""
            )
            line_hint = (
                line_match.group(1).strip().splitlines()[0].strip()
                if line_match
                else ""
            )

            # Try to extract a file/line hint from the Code: field as fallback.
            code_hint = ""
            code_match = re.search(
                r"Code\s*:\s*(.+?)(?=\s*(?:Reason|Suggestion)\s*:|$)",
                raw_block,
                re.DOTALL | re.IGNORECASE
            )
            if code_match:
                # Take only the first line of the code snippet as the hint.
                code_hint = (
                    code_match.group(1).strip().splitlines()[0].strip()
                )

            if not file_hint or not line_hint:
                fallback_file, fallback_line = _extract_file_and_line(code_hint)
                if not file_hint:
                    file_hint = fallback_file
                if not line_hint:
                    line_hint = fallback_line

            issues.append({
                "file": file_hint,
                "line": line_hint,
                "issue": issue_text,
            })

    else:
        # Fallback: numbered list items (e.g. '1. Something wrong').
        fallback_blocks = re.split(
            r"\n\s*\d+[\.\)]\s+",
            review_text.strip()
        )
        for block in fallback_blocks:
            block = block.strip()
            if not block:
                continue
            issues.append({
                "file": "",
                "line": "",
                "issue": re.sub(r"\s+", " ", block).strip(),
            })

    return issues


def _extract_file_and_line(code_hint: str) -> tuple[str, str]:
    """
    Attempt to extract a file path and line number from a code snippet hint.

    Handles patterns like:
        src/utils.py:42  ->  file="src/utils.py", line="42"
        utils.py line 10 ->  file="utils.py",     line="10"
        src/utils.py     ->  file="src/utils.py", line=""
    """
    if not code_hint:
        return "", ""

    # Pattern: path/to/file.py:42
    match = re.match(
        r"^(?P<path>[^\s:]+\.\w+)\s*:\s*(?P<line>\d+)",
        code_hint
    )
    if match:
        return match.group("path"), match.group("line")

    # Pattern: path/to/file.py line 42
    match = re.match(
        r"^(?P<path>[^\s]+\.\w+)\s+line\s+(?P<line>\d+)",
        code_hint,
        re.IGNORECASE
    )
    if match:
        return match.group("path"), match.group("line")

    # Pattern: bare file path with a known extension
    match = re.match(
        r"^(?P<path>[^\s]+\.(?:py|js|ts|tsx|jsx|java|go|cs|rb|php|swift|kt|rs))",
        code_hint,
        re.IGNORECASE
    )
    if match:
        return match.group("path"), ""

    return "", ""


# ---------------------------------------------------------------------------
# Public: generate_word_report
# ---------------------------------------------------------------------------

_REPORT_CACHE: dict[str, bytes] = {}


def generate_word_report(
    repo_name: str,
    pr_number: int,
    review_text: str,
    report_dir: str
) -> str:
    """
    Generate a .docx report containing only the AI review issues.

    Parameters
    ----------
    repo_name   : Full repository name, e.g. "org/repo"
    pr_number   : Pull Request number
    review_text : Raw AI review string returned by review_code()
    report_dir  : Directory to save the report in (created if absent)

    Returns
    -------
    Absolute path to the generated .docx file.
    """
    os.makedirs(report_dir, exist_ok=True)

    issues = parse_review_issues(review_text)
    review_date = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    # Build a safe filename slug from the repo name.
    repo_slug = re.sub(r"[^a-zA-Z0-9_-]", "_", repo_name)
    timestamp = datetime.datetime.utcnow().strftime("%Y%m%d%H%M%S")
    filename = f"pr_{repo_slug}_{pr_number}_{timestamp}.docx"
    file_path = os.path.join(report_dir, filename)

    doc = _build_document(repo_name, pr_number, review_date, issues)
    doc.save(file_path)

    try:
        import io
        buf = io.BytesIO()
        doc.save(buf)
        _REPORT_CACHE[filename] = buf.getvalue()
    except Exception as cache_err:
        print(f"[REPORT CACHE WARNING] {cache_err}")

    print(f"[REPORT] Word report saved: {file_path}")
    return file_path


def generate_mismatch_report(
    repo_name: str,
    pr_number: int,
    mismatches: list[dict],
    report_dir: str
) -> str:
    """
    Generate a .docx report from structured comment-validation mismatches.

    Each mismatch dict is expected to have: file, line, comment, reason.
    File and Line fields are populated directly from the mismatch data.

    Parameters
    ----------
    repo_name   : Full repository name, e.g. "org/repo"
    pr_number   : Pull Request number
    mismatches  : List of mismatch dicts from validate_code_comments()
    report_dir  : Directory to save the report in (created if absent)

    Returns
    -------
    Absolute path to the generated .docx file.
    """
    os.makedirs(report_dir, exist_ok=True)

    # Convert structured mismatches into the common issue dict format.
    # The mismatch reason becomes the issue description; file and line
    # are carried through directly so they appear in the report.
    issues = [
        {
            "file": item.get("file", "").strip(),
            "line": item.get("line", "").strip(),
            "issue": item.get("reason", "").strip(),
        }
        for item in mismatches
        if item.get("reason", "").strip()
    ]

    review_date = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    repo_slug = re.sub(r"[^a-zA-Z0-9_-]", "_", repo_name)
    timestamp = datetime.datetime.utcnow().strftime("%Y%m%d%H%M%S")
    filename = (
        f"comment_validation_{repo_slug}_{pr_number}_{timestamp}.docx"
    )
    file_path = os.path.join(report_dir, filename)

    doc = _build_document(repo_name, pr_number, review_date, issues)
    doc.save(file_path)

    try:
        import io
        buf = io.BytesIO()
        doc.save(buf)
        _REPORT_CACHE[filename] = buf.getvalue()
    except Exception as cache_err:
        print(f"[REPORT CACHE WARNING] {cache_err}")

    print(f"[REPORT] Comment validation report saved: {file_path}")
    return file_path


def get_report_bytes(filename: str, report_dir: str) -> bytes | None:
    """
    Retrieve report file bytes by filename.
    Checks disk first, then memory cache, then regenerates from GitHub on demand.
    """
    file_path = os.path.join(report_dir, filename)

    if os.path.isfile(file_path):
        try:
            with open(file_path, "rb") as f:
                return f.read()
        except Exception:
            pass

    if filename in _REPORT_CACHE:
        return _REPORT_CACHE[filename]

    return _regenerate_report_on_demand(filename, report_dir)


def _regenerate_report_on_demand(filename: str, report_dir: str) -> bytes | None:
    """
    Regenerate report file on demand if server restarted and disk was wiped.
    Fetches the posted PR comments from GitHub API and rebuilds the .docx.
    """
    pattern = (
        r"^(?P<prefix>pr|comment_validation)_(?P<repo_slug>.+?)_"
        r"(?P<pr_number>\d+)_(?P<timestamp>\d+)\.docx$"
    )
    match = re.match(pattern, filename, re.IGNORECASE)
    if not match:
        return None

    prefix = match.group("prefix").lower()
    repo_slug = match.group("repo_slug")
    pr_number = int(match.group("pr_number"))

    if "__" in repo_slug:
        repo_name = repo_slug.replace("__", "/")
    else:
        repo_name = repo_slug.replace("_", "/", 1)

    print(
        f"[REPORT REGEN] Regenerating report on demand for {repo_name} PR #{pr_number} ({filename})"
    )

    try:
        import io
        import requests
        from github.github_auth import generate_installation_token

        token = generate_installation_token()
        url = f"https://api.github.com/repos/{repo_name}/issues/{pr_number}/comments"
        headers = {
            "Authorization": f"Bearer {token}",
            "Accept": "application/vnd.github+json"
        }
        res = requests.get(url, headers=headers, params={"per_page": 100})
        if res.status_code != 200:
            print(f"[REPORT REGEN] GitHub API HTTP {res.status_code}: {res.text}")
            return None

        comments = res.json()
        target_comment = None

        if prefix == "pr":
            for c in comments:
                body = c.get("body") or ""
                if "## AI Code Review" in body or "Issue:" in body:
                    target_comment = body
                    break
            if target_comment:
                clean_body = re.sub(
                    r"\n\n---\n📄 \[Download Word Report\].*",
                    "",
                    target_comment,
                    flags=re.DOTALL
                )
                clean_body = re.sub(r"^## AI Code Review\n\n", "", clean_body)
                issues = parse_review_issues(clean_body)
                review_date = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
                doc = _build_document(repo_name, pr_number, review_date, issues)
                os.makedirs(report_dir, exist_ok=True)
                file_path = os.path.join(report_dir, filename)
                doc.save(file_path)
                buf = io.BytesIO()
                doc.save(buf)
                data = buf.getvalue()
                _REPORT_CACHE[filename] = data
                return data

        elif prefix == "comment_validation":
            for c in comments:
                body = c.get("body") or ""
                if "Comment Validation Failed" in body:
                    target_comment = body
                    break
            if target_comment:
                pattern_mismatch = re.compile(
                    r"File:\s*(?P<file>.+?)\s*\n"
                    r"Line:\s*(?P<line>.+?)\s*\n"
                    r'Comment:\s*"(?P<comment>.+?)"\s*\n'
                    r"Code:\s*(?P<code>.+?)\s*\n"
                    r"Reason:\s*(?P<reason>.+?)(?=\n---|\nFile:|\n\nPlease|\Z)",
                    re.DOTALL | re.IGNORECASE
                )
                mismatches = []
                for m in pattern_mismatch.finditer(target_comment):
                    mismatches.append({
                        "file": m.group("file").strip(),
                        "line": m.group("line").strip(),
                        "comment": m.group("comment").strip(),
                        "code": m.group("code").strip(),
                        "reason": m.group("reason").strip()
                    })

                if mismatches:
                    issues = [
                        {
                            "file": item.get("file", "").strip(),
                            "line": item.get("line", "").strip(),
                            "issue": item.get("reason", "").strip(),
                        }
                        for item in mismatches
                        if item.get("reason", "").strip()
                    ]
                    review_date = datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
                    doc = _build_document(repo_name, pr_number, review_date, issues)
                    os.makedirs(report_dir, exist_ok=True)
                    file_path = os.path.join(report_dir, filename)
                    doc.save(file_path)
                    buf = io.BytesIO()
                    doc.save(buf)
                    data = buf.getvalue()
                    _REPORT_CACHE[filename] = data
                    return data

    except Exception as err:
        print(f"[REPORT REGEN ERROR] {err}")

    return None



# ---------------------------------------------------------------------------
# Internal: document builder
# ---------------------------------------------------------------------------

def _build_document(
    repo_name: str,
    pr_number: int,
    review_date: str,
    issues: list[dict]
) -> Document:
    """Build and return the python-docx Document object."""
    doc = Document()

    # ----- Document title -----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AI Pull Request Review Report")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue

    doc.add_paragraph()  # spacer

    # ----- Metadata block -----
    _add_meta_line(doc, "Repository", repo_name)
    _add_meta_line(doc, "Pull Request", f"#{pr_number}")
    _add_meta_line(doc, "Review Date", review_date)

    doc.add_paragraph()  # spacer

    # ----- Issues -----
    if not issues:
        _add_separator(doc)
        no_issue_para = doc.add_paragraph()
        no_issue_run = no_issue_para.add_run(
            "No issues were detected in this Pull Request."
        )
        no_issue_run.italic = True
        _add_separator(doc)
    else:
        for idx, item in enumerate(issues, start=1):
            _add_separator(doc)

            # Issue heading
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(f"Issue {idx}")
            heading_run.bold = True
            heading_run.font.size = Pt(12)

            # File (only when the information is available)
            if item.get("file"):
                _add_issue_field(doc, "File", item["file"])

            # Line (only when the information is available)
            if item.get("line"):
                _add_issue_field(doc, "Line", item["line"])

            # Issue description
            _add_issue_field(doc, "Issue", item["issue"])

        _add_separator(doc)

    # ----- Footer: total issues -----
    doc.add_paragraph()  # spacer
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(f"Total Issues: {len(issues)}")
    footer_run.bold = True
    footer_run.font.size = Pt(11)

    return doc


def _add_meta_line(doc: Document, label: str, value: str) -> None:
    """Add a bold-label + normal-value line to the document."""
    para = doc.add_paragraph()
    label_run = para.add_run(f"{label}: ")
    label_run.bold = True
    para.add_run(value)


def _add_separator(doc: Document) -> None:
    """Add a horizontal separator paragraph."""
    sep_para = doc.add_paragraph(_SEPARATOR)
    sep_para.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def _add_issue_field(doc: Document, label: str, value: str) -> None:
    """Add a 'Label: value' line inside an issue block."""
    para = doc.add_paragraph()
    label_run = para.add_run(f"{label}: ")
    label_run.bold = True
    para.add_run(value)
